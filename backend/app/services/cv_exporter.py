"""
cv_exporter.py
Exports CV to PDF (via WeasyPrint) and DOCX (via python-docx).
"""
import io
from weasyprint import HTML as WeasyHTML
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_pdf(html: str) -> bytes:
    """Convert an HTML string to PDF bytes using WeasyPrint."""
    pdf_bytes = WeasyHTML(string=html).write_pdf()
    return pdf_bytes


def export_docx(profile_json: dict, schema_json: dict) -> bytes:
    """
    Generate a simplified DOCX from profile_json.
    Layout is single-column for maximum compatibility.
    """
    doc = Document()

    # Remove default margins slightly
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    colors = schema_json.get("colors", {})
    accent_hex = colors.get("accent", "4361EE").lstrip("#")
    try:
        accent_rgb = tuple(int(accent_hex[i:i+2], 16) for i in (0, 2, 4))
    except Exception:
        accent_rgb = (67, 97, 238)

    basic = profile_json.get("basic_info", {})

    # --- Header: Name ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(basic.get("full_name", "No Name"))
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = RGBColor(*accent_rgb)

    # Contact line
    contact_parts = []
    if basic.get("email"):
        contact_parts.append(basic["email"])
    if basic.get("phone"):
        contact_parts.append(basic["phone"])
    if basic.get("linkedin"):
        contact_parts.append(basic["linkedin"])

    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(10)

    # Summary
    if basic.get("summary"):
        doc.add_paragraph()
        _add_section_heading(doc, "Summary", accent_rgb)
        doc.add_paragraph(basic["summary"])

    # Work Experiences
    work_exps = profile_json.get("work_experiences", [])
    if work_exps:
        _add_section_heading(doc, "Work Experience", accent_rgb)
        for exp in work_exps:
            role_para = doc.add_paragraph()
            role_run = role_para.add_run(exp.get("role", "") + " @ " + exp.get("company", ""))
            role_run.bold = True
            role_run.font.size = Pt(11)

            date_range = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
            date_para = doc.add_paragraph(date_range)
            for run in date_para.runs:
                run.italic = True
                run.font.size = Pt(10)

            if exp.get("description"):
                doc.add_paragraph(exp["description"])

    # Educations
    educations = profile_json.get("educations", [])
    if educations:
        _add_section_heading(doc, "Education", accent_rgb)
        for edu in educations:
            edu_para = doc.add_paragraph()
            edu_run = edu_para.add_run(edu.get("degree", "") + " - " + edu.get("school", ""))
            edu_run.bold = True
            edu_run.font.size = Pt(11)
            if edu.get("major"):
                doc.add_paragraph(edu["major"])

    # Skills
    skills = profile_json.get("skills", [])
    if skills:
        _add_section_heading(doc, "Skills", accent_rgb)
        for skill_group in skills:
            category = skill_group.get("category", "")
            items = skill_group.get("items", [])
            if items:
                skill_text = (f"{category}: " if category else "") + ", ".join(items)
                doc.add_paragraph(skill_text)

    # Projects
    projects = profile_json.get("projects", [])
    if projects:
        _add_section_heading(doc, "Projects", accent_rgb)
        for project in projects:
            proj_para = doc.add_paragraph()
            proj_run = proj_para.add_run(project.get("name", ""))
            proj_run.bold = True
            if project.get("description"):
                doc.add_paragraph(project["description"])
            if project.get("tech_stack"):
                doc.add_paragraph("Tech: " + ", ".join(project["tech_stack"]))

    # Certifications
    certs = profile_json.get("certifications", [])
    if certs:
        _add_section_heading(doc, "Certifications", accent_rgb)
        for cert in certs:
            cert_text = cert.get("name", "")
            if cert.get("issuer"):
                cert_text += f" ({cert['issuer']})"
            doc.add_paragraph(cert_text)

    # Languages
    langs = profile_json.get("languages", [])
    if langs:
        _add_section_heading(doc, "Languages", accent_rgb)
        for lang in langs:
            doc.add_paragraph(f"{lang.get('language', '')} - {lang.get('level', '')}")

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_section_heading(doc: Document, title: str, accent_rgb: tuple):
    """Add a styled section heading with an underline-like separator."""
    heading_para = doc.add_paragraph()
    heading_run = heading_para.add_run(title.upper())
    heading_run.bold = True
    heading_run.font.size = Pt(12)
    heading_run.font.color.rgb = RGBColor(*accent_rgb)
    # Horizontal rule via bottom border on paragraph
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = heading_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*accent_rgb))
    pBdr.append(bottom)
    pPr.append(pBdr)
